from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(r"E:\石化厂赛事项目\石化厂文件\过程中产生的技术文档\文档.docx")
OUTPUT = Path(r"E:\石化厂赛事项目\石化厂文件\过程中产生的技术文档\文档_优化版.docx")


def set_cell_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_code_block(document: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(12)
        paragraph.paragraph_format.right_indent = Pt(12)
        set_cell_shading(paragraph, "F6F8FA")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(31, 35, 40)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(3)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.color.rgb = RGBColor(31, 35, 40)


def build_document() -> Document:
    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("纯前端 / UI 设计中的技术问题及解决办法")
    title_run.bold = True
    title_run.font.name = "微软雅黑"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title_run.font.size = Pt(18)

    add_heading(document, "一、背景", 1)
    add_body(
        document,
        "石安盾平台功能模块较多，需要通过左侧菜单为用户提供稳定、清晰的导航入口。"
        "在完整展开状态下，菜单会占用较多横向空间，压缩右侧正文区域；因此前端采用可折叠侧边栏，"
        "通过 toggle 按钮在“完整菜单”和“窄图标菜单”之间切换。",
    )
    add_body(
        document,
        "该方案的核心难点不在于实现收起状态本身，而在于让展开与收起之间的过渡足够自然："
        "图标位置应保持稳定，文字与箭头应平滑淡出，侧边栏边界应连续收缩，避免出现跳动、闪烁或布局突变。"
    )
    add_body(document, "图 1-1：LeftBar 展开与收起状态对比（示意）")

    add_heading(document, "二、问题现象", 1)
    add_body(
        document,
        "点击“收起菜单”时，理想效果是左侧图标保持原位不动，右边框从右向左平滑收缩，"
        "文字和子菜单的下三角图标逐渐淡出。实际实现中，文字与下三角图标会在点击瞬间消失，"
        "左侧图标也会突然跳动，随后侧边栏宽度才开始执行 transition。最终视觉表现像是“先闪一下，再播放动画”。",
    )

    add_heading(document, "三、原因分析", 1)
    add_body(
        document,
        "根本原因是状态变量 isCollapsed 改变后，React 先触发了 DOM 结构和内部布局的同步变化，"
        "随后外层容器的 width 才进入 180ms 的 CSS 过渡。也就是说，真正需要动画的元素在动画开始前已经被移除或重新排版。"
    )
    add_body(document, "原始实现中，文字和箭头采用条件渲染：")
    add_code_block(
        document,
        """
{!isCollapsed && <Word>{item.label}</Word>}
{!isCollapsed && (
  <ChevronIcon $isOpen={isOpen}>
    <Icon id="ChevronDown" size={15} strokeWidth={2} />
  </ChevronIcon>
)}
""",
    )
    add_body(document, "这会造成三个直接后果：")
    add_bullet(document, "文字和下三角图标在点击瞬间从 DOM 中卸载，没有淡出过程。")
    add_bullet(document, "菜单项内部布局重新计算，图标从原来的左侧位置跳到收起态位置。")
    add_bullet(document, "外层侧边栏 width 仍在过渡，但内部元素已经完成突变，导致视觉上不连贯。")

    add_heading(document, "四、修复方案", 1)
    add_body(
        document,
        "修复原则是：图标所在的布局基准必须在展开和收起两种状态下保持一致。"
        "具体做法是使用 CSS Grid 固定左侧图标列，文字列和箭头只做视觉收缩或淡出，不再通过 React 条件渲染直接卸载。"
    )

    add_heading(document, "1. 固定菜单项的网格布局", 2)
    add_body(document, "菜单项统一采用两列 grid：第一列固定放图标，第二列放文字。收起时只压缩第二列，不改变第一列。")
    add_code_block(
        document,
        """
const itemStyles = `
  position: relative;
  width: 100%;
  display: grid;
  grid-template-columns: 18px minmax(0, 1fr);
  column-gap: 10px;
  padding: 8px 30px 8px 12px;
  border-radius: 4px;
  color: inherit;
  font-size: ${13 / 16}rem;
  text-decoration: none;
`;

const collapsedItemStyles = `
  grid-template-columns: 18px 0;
  column-gap: 0;
  padding-right: 12px;
`;
""",
    )

    add_heading(document, "2. 保留文字节点，用 opacity 控制淡出", 2)
    add_body(document, "文字始终保留在第二列，不再根据 isCollapsed 卸载。收起时只改变透明度，让浏览器负责动画。")
    add_code_block(
        document,
        """
<Icon id={item.icon} size={18} />
<Word $isCollapsed={isCollapsed}>{item.label}</Word>

const Word = styled.p`
  grid-column: 2;
  min-width: 0;
  overflow: hidden;
  opacity: ${(p) => (p.$isCollapsed ? 0 : 1)};
  white-space: nowrap;
  text-align: left;
  transform: translateY(-1px);
  transition: opacity 120ms ease;
`;
""",
    )

    add_heading(document, "3. 将下三角图标从布局流中移出", 2)
    add_body(
        document,
        "下三角图标如果继续参与 grid 布局，在收起时可能被压缩列挤到左侧图标附近。"
        "因此将其改为绝对定位，固定在菜单项右侧，仅做透明度和旋转过渡。"
    )
    add_code_block(
        document,
        """
<ChevronIcon $isOpen={isOpen} $isCollapsed={isCollapsed}>
  <Icon id="ChevronDown" size={15} strokeWidth={2} />
</ChevronIcon>

const ChevronIcon = styled.div`
  position: absolute;
  right: 8px;
  top: 50%;
  opacity: ${(p) => (p.$isCollapsed ? 0 : 1)};
  pointer-events: none;
  transform: translateY(-50%)
    rotate(${(p) => (p.$isOpen ? "180deg" : "0deg")});
  transition: opacity 120ms ease, transform 160ms ease;
`;
""",
    )

    add_heading(document, "4. 外层侧边栏只负责宽度过渡", 2)
    add_body(document, "外层 Wrapper 保留 width transition。由于内部图标列位置固定，收起时视觉上会更加稳定。")
    add_code_block(
        document,
        """
const Wrapper = styled.div`
  width: ${(p) => (p.$isCollapsed ? "4rem" : "10rem")};
  padding: 16px 8px 0 8px;
  display: flex;
  flex-direction: column;
  gap: 12px;
  border-right: 1px ${COLORS.gray90} solid;
  transition: width 180ms ease;
`;
""",
    )

    add_heading(document, "五、最终效果", 1)
    add_body(document, "修改后，点击“收起菜单”时的执行过程变为：")
    add_bullet(document, "侧边栏宽度开始从 10rem 平滑缩小到 4rem。")
    add_bullet(document, "文字和下三角图标通过 opacity 逐渐淡出。")
    add_bullet(document, "左侧图标始终固定在同一条 grid 图标列上，不再发生横向跳动。")
    add_bullet(document, "下三角图标固定在右侧淡出，不会在收起瞬间闪到左侧图标位置。")

    add_heading(document, "六、经验总结", 1)
    add_bullet(document, "UI 过渡的关键不是简单添加 transition，而是避免动画过程中 DOM 结构和布局规则突然变化。")
    add_bullet(document, "需要稳定的元素，例如左侧图标，应使用固定列或固定定位保持位置不变。")
    add_bullet(document, "需要消失的元素，例如文字和箭头，应优先使用 opacity、width、max-width 或 grid 列宽收缩处理。")
    add_bullet(document, "不要在收起态突然改用 justify-content: center，否则图标会因对齐方式变化而跳动。")
    add_bullet(document, "React 状态适合描述组件状态，具体过渡细节应尽量交给 CSS 完成。")

    return document


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
