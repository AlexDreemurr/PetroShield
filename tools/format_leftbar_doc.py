from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOC_PATH = Path(r"E:\石化厂赛事项目\石化厂文件\过程中产生的技术文档\文档_优化版.docx")
OUTPUT_PATH = Path(r"E:\石化厂赛事项目\石化厂文件\过程中产生的技术文档\文档_优化版_配图小字号.docx")
IMAGE_BEFORE = Path(r"E:\石化厂赛事项目\石化厂文件\过程中产生的技术文档\LeftBar收缩前.png")
IMAGE_AFTER = Path(r"E:\石化厂赛事项目\石化厂文件\过程中产生的技术文档\LeftBar收缩后.png")

FONT_NAME = "宋体"
HEADING_SIZE = Pt(9)  # 小五
BODY_SIZE = Pt(8)
CODE_SIZE = Pt(7)


def set_run_font(run, size=BODY_SIZE, font_name=FONT_NAME, bold=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = size
    if bold is not None:
        run.bold = bold


def paragraph_has_shading(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    return p_pr.find(qn("w:shd")) is not None


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def is_heading_text(text: str, index: int) -> bool:
    if index == 0:
        return True
    if text.startswith(("一、", "二、", "三、", "四、", "五、", "六、")):
        return True
    if text.startswith(("1. ", "2. ", "3. ", "4. ")):
        return True
    return False


def looks_like_code(text: str, paragraph) -> bool:
    if paragraph_has_shading(paragraph):
        return True
    stripped = text.strip()
    if not stripped:
        return False
    code_markers = (
        "const ",
        "display:",
        "grid-template-columns",
        "column-gap",
        "padding:",
        "border-radius",
        "font-size:",
        "text-decoration",
        "opacity:",
        "transition:",
        "transform:",
        "position:",
        "right:",
        "top:",
        "pointer-events",
        "width:",
        "gap:",
        "border-right",
        "{!isCollapsed",
        "<Word",
        "<ChevronIcon",
        "<Icon",
        "</ChevronIcon",
        "`;",
        "};",
    )
    return stripped.startswith(code_markers) or stripped in {"`", ")", "}"}


def insert_picture_table_after(paragraph, document: Document) -> None:
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    labels = ["图 1-1 LeftBar 展开状态", "图 1-2 LeftBar 收缩状态"]
    images = [IMAGE_BEFORE, IMAGE_AFTER]
    widths = [Inches(1.35), Inches(0.55)]

    for col, label in enumerate(labels):
        cell = table.cell(0, col)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_run_font(run, BODY_SIZE)

    for col, image_path in enumerate(images):
        cell = table.cell(1, col)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=widths[col])

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "nil")
                tc_pr.append(border)

    table_parent = table._tbl.getparent()
    if table_parent is not None:
        table_parent.remove(table._tbl)
    paragraph._p.addnext(table._tbl)


def apply_document_format(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = BODY_SIZE

    for i, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        code = looks_like_code(text, paragraph)
        heading = is_heading_text(text, i)

        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05

        if i == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if code:
            set_paragraph_shading(paragraph, "F6F8FA")
            paragraph.paragraph_format.left_indent = Pt(10)
            paragraph.paragraph_format.right_indent = Pt(10)
            paragraph.paragraph_format.space_after = Pt(0)

        for run in paragraph.runs:
            if code:
                set_run_font(run, CODE_SIZE, FONT_NAME)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
                run.font.color.rgb = RGBColor(31, 35, 40)
            elif heading:
                set_run_font(run, HEADING_SIZE, FONT_NAME, bold=True)
            else:
                set_run_font(run, BODY_SIZE, FONT_NAME)


def main() -> None:
    for path in [DOC_PATH, IMAGE_BEFORE, IMAGE_AFTER]:
        if not path.exists():
            raise FileNotFoundError(path)

    document = Document(DOC_PATH)

    if len(document.inline_shapes) == 0:
        for paragraph in document.paragraphs:
            if "图 1-1" in paragraph.text:
                insert_picture_table_after(paragraph, document)
                break

    apply_document_format(document)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
